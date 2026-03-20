"""
文件操作能力工具。

通过 contextvars 获取 FileService 和 user_id，
Agent 可读取用户上传的文件、列出文件、分析文件结构。

所有工具为 read_only=True（不修改文件系统）。
"""

from __future__ import annotations

import logging

from core.context import current_event_bus, current_user_id, current_tenant_id, current_session_id
from core.tool_registry import ToolRegistry

logger = logging.getLogger(__name__)

file_capability_registry = ToolRegistry()


def _get_file_service():
    """从 contextvars 获取 FileService。"""
    from core.context import current_file_service
    service = current_file_service.get()
    if service is None:
        raise RuntimeError("FileService not available (not injected)")
    return service


@file_capability_registry.tool(
    description=(
        "读取当前会话中用户上传的文件内容。"
        "传入 file_id (从 list_user_files 获取)，返回文件的提取文本。"
        "只能读取当前会话的文件，不能跨会话访问。"
        "支持 PDF/DOCX/TXT/CSV/JSON 等格式的文本提取。"
        "图片文件返回尺寸等元信息。"
        "大文件自动分页 — 使用 offset/limit 参数分段读取。"
    ),
    read_only=True,
)
def read_uploaded_file(
    file_id: str,           # 文件 ID
    offset: int = 0,        # 起始字符偏移 (默认从头)
    limit: int = 0,         # 读取字符数 (0 = 使用默认页大小)
) -> dict:
    """读取用户上传的文件，返回提取的文本内容 (支持分页)。"""
    from config import settings

    service = _get_file_service()
    tenant_id = current_tenant_id.get()
    user_id = current_user_id.get()
    session_id = current_session_id.get("")

    try:
        metadata, _ = service.get_file(tenant_id, user_id, file_id)

        # 会话隔离: 只能读取当前会话的文件
        if session_id and metadata.session_id and metadata.session_id != session_id:
            return {"error": f"File {file_id} does not belong to current session"}

        text = service.extract_text(tenant_id, user_id, file_id)

        from core.text_utils import paginate_text
        page = paginate_text(text, offset=offset, limit=limit,
                             context_window=settings.agent_model_context_window)

        result = {
            "file_id": file_id,
            "filename": metadata.filename,
            "content_type": metadata.content_type,
            "size_bytes": metadata.size_bytes,
            "text": page.text,
        }
        if page.has_more or page.offset > 0 or page.length < page.total_chars:
            result["pagination"] = {
                "offset": page.offset,
                "length": page.length,
                "total_chars": page.total_chars,
                "has_more": page.has_more,
                "next_offset": page.next_offset,
            }
            if page.has_more:
                result["hint"] = (
                    f"文件共 {page.total_chars} 字符，当前显示 {page.offset}-{page.next_offset}。"
                    f"使用 offset={page.next_offset} 继续读取下一段。"
                )
        return result
    except FileNotFoundError:
        return {"error": f"File not found: {file_id}"}
    except Exception as e:
        logger.error(f"read_uploaded_file error: {e}")
        return {"error": str(e)}


@file_capability_registry.tool(
    description=(
        "列出当前会话中用户上传的文件。"
        "返回文件列表，包含 file_id/filename/content_type/size_bytes。"
        "只列出当前会话关联的文件，不能跨会话访问。"
        "使用 file_id 可调用 read_uploaded_file 读取文件内容。"
    ),
    read_only=True,
)
def list_user_files() -> dict:
    """列出当前会话关联的文件。"""
    service = _get_file_service()
    tenant_id = current_tenant_id.get()
    user_id = current_user_id.get()
    session_id = current_session_id.get("")

    try:
        if session_id:
            files = service.list_files_by_session(tenant_id, user_id, session_id)
        else:
            files = service.list_files(tenant_id, user_id)
        return {
            "user_id": user_id,
            "session_id": session_id,
            "file_count": len(files),
            "files": [
                {
                    "file_id": f.file_id,
                    "filename": f.filename,
                    "content_type": f.content_type,
                    "size_bytes": f.size_bytes,
                }
                for f in files
            ],
        }
    except Exception as e:
        logger.error(f"list_user_files error: {e}")
        return {"error": str(e)}


@file_capability_registry.tool(
    description=(
        "按需读取知识库中的文件内容。"
        "传入 file_id (从系统提示中的 <knowledge> 索引获取)，返回文件的提取文本。"
        "仅在需要引用知识库中某个文件的详细内容时调用，不要预先读取所有文件。"
    ),
    read_only=True,
)
def read_knowledge_file(
    file_id: str,           # 知识库文件 ID
    offset: int = 0,        # 起始字符偏移 (默认从头)
    limit: int = 0,         # 读取字符数 (0 = 使用默认页大小)
) -> dict:
    """按需读取知识库文件，返回提取的文本内容 (支持分页)。"""
    from config import settings

    try:
        from dependencies import get_knowledge_service
        kb_service = get_knowledge_service()
    except Exception as e:
        return {"error": f"Knowledge service not available: {e}"}

    try:
        text = kb_service.extract_text(file_id)
        meta = kb_service.get_file_meta(file_id)
        if not meta:
            return {"error": f"Knowledge file not found: {file_id}"}

        from core.text_utils import paginate_text
        page = paginate_text(text, offset=offset, limit=limit,
                             context_window=settings.agent_model_context_window)

        result = {
            "file_id": file_id,
            "filename": meta.filename,
            "description": meta.description,
            "text": page.text,
        }
        if page.has_more or page.offset > 0 or page.length < page.total_chars:
            result["pagination"] = {
                "offset": page.offset,
                "length": page.length,
                "total_chars": page.total_chars,
                "has_more": page.has_more,
                "next_offset": page.next_offset,
            }
            if page.has_more:
                result["hint"] = (
                    f"文件共 {page.total_chars} 字符，当前显示 {page.offset}-{page.next_offset}。"
                    f"使用 offset={page.next_offset} 继续读取下一段。"
                )
        return result
    except Exception as e:
        logger.error(f"read_knowledge_file error: {e}")
        return {"error": str(e)}


@file_capability_registry.tool(
    description=(
        "将内容添加到知识库。"
        "可以传入文本内容 (text_content + filename)，或传入已上传文件的 file_id 将其复制到知识库。"
        "文件会自动创建元数据并更新知识库索引。"
        "scope 始终为 user (用户个人知识库)，普通用户不可写入 global。"
    ),
    read_only=False,
)
def add_to_knowledge_base(
    filename: str,              # 文件名 (如 "report.md")
    description: str = "",      # 文件描述
    text_content: str = "",     # 文本内容 (与 source_file_id 二选一)
    source_file_id: str = "",   # 从已上传文件复制 (与 text_content 二选一)
) -> dict:
    """将内容添加到用户知识库，自动创建元数据和索引。"""
    tenant_id = current_tenant_id.get()
    user_id = current_user_id.get()

    if not text_content and not source_file_id:
        return {"error": "必须提供 text_content 或 source_file_id"}

    try:
        from dependencies import get_knowledge_service
        kb_service = get_knowledge_service()
    except Exception as e:
        return {"error": f"Knowledge service not available: {e}"}

    try:
        # 获取文件内容
        if source_file_id:
            # 从已上传文件复制
            file_service = _get_file_service()
            try:
                metadata, content = file_service.get_file(tenant_id, user_id, source_file_id)
                if not filename:
                    filename = metadata.filename
            except FileNotFoundError:
                return {"error": f"Source file not found: {source_file_id}"}
        else:
            content = text_content.encode("utf-8")

        # 强制 user scope — 普通用户不可写 global
        scope = "user"

        meta = kb_service.add_file(
            tenant_id=tenant_id,
            user_id=user_id,
            filename=filename,
            content=content,
            scope=scope,
            description=description,
        )

        # 触发异步索引 (后台生成摘要 + 更新 _index.md)
        bus = current_event_bus.get()
        if bus:
            import asyncio
            from api.knowledge_routes import _auto_index
            try:
                loop = asyncio.get_event_loop()
                loop.create_task(_auto_index(kb_service, meta, tenant_id, user_id))
            except RuntimeError:
                logger.debug("No event loop for auto-index, skipping")

        from dataclasses import asdict
        result = asdict(meta)
        result["status"] = "ok"
        result["message"] = f"文件 '{meta.filename}' 已添加到个人知识库"
        return result

    except Exception as e:
        logger.error(f"add_to_knowledge_base error: {e}")
        return {"error": str(e)}


@file_capability_registry.tool(
    description=(
        "分析用户上传文件的结构和基本信息。"
        "返回文件的详细元信息：大小、类型、页数(PDF)、行数(文本)等。"
    ),
    read_only=True,
)
def analyze_file(file_id: str) -> dict:  # 文件 ID
    """分析文件结构，返回详细元信息。"""
    service = _get_file_service()
    tenant_id = current_tenant_id.get()
    user_id = current_user_id.get()

    try:
        metadata, content = service.get_file(tenant_id, user_id, file_id)
        ext = metadata.filename.rsplit(".", 1)[-1].lower() if "." in metadata.filename else ""

        analysis = {
            "file_id": file_id,
            "filename": metadata.filename,
            "content_type": metadata.content_type,
            "size_bytes": metadata.size_bytes,
            "sha256": metadata.sha256,
        }

        # 按类型补充分析
        if ext == "pdf":
            try:
                import io
                from PyPDF2 import PdfReader
                reader = PdfReader(io.BytesIO(content))
                analysis["page_count"] = len(reader.pages)
                analysis["format"] = "PDF"
            except Exception:
                analysis["format"] = "PDF (parse error)"

        elif ext in ("docx", "doc"):
            try:
                import io
                from docx import Document
                doc = Document(io.BytesIO(content))
                analysis["paragraph_count"] = len(doc.paragraphs)
                analysis["table_count"] = len(doc.tables)
                analysis["format"] = "DOCX"
            except Exception:
                analysis["format"] = "DOCX (parse error)"

        elif ext in ("txt", "csv", "json", "xml", "yaml", "yml", "md"):
            try:
                text = content.decode("utf-8", errors="replace")
                lines = text.split("\n")
                analysis["line_count"] = len(lines)
                analysis["char_count"] = len(text)
                analysis["format"] = ext.upper()
            except Exception:
                analysis["format"] = f"{ext.upper()} (parse error)"

        elif ext in ("png", "jpg", "jpeg", "gif", "bmp", "webp"):
            try:
                import io
                from PIL import Image
                img = Image.open(io.BytesIO(content))
                analysis["width"] = img.width
                analysis["height"] = img.height
                analysis["image_mode"] = img.mode
                analysis["image_format"] = img.format
                analysis["format"] = "Image"
            except Exception:
                analysis["format"] = "Image (parse error)"
        else:
            analysis["format"] = ext.upper() or "Unknown"

        return analysis

    except FileNotFoundError:
        return {"error": f"File not found: {file_id}"}
    except Exception as e:
        logger.error(f"analyze_file error: {e}")
        return {"error": str(e)}


@file_capability_registry.tool(
    description=(
        "处理文件的多模态内容 (A4: 4i)。"
        "图片文件: 转 base64 (自动压缩大图)。"
        "PDF文件: 提取文本内容 (支持多页)。"
        "文本文件: 直接返回内容。"
        "返回结构包含 content_type 和对应的数据。"
    ),
    read_only=True,
)
def process_file_content(file_id: str) -> dict:  # 文件 ID
    """处理文件内容，返回适合 LLM 消费的格式。"""
    from services.content_processor import process_file

    service = _get_file_service()
    tenant_id = current_tenant_id.get()
    user_id = current_user_id.get()

    try:
        metadata, content = service.get_file(tenant_id, user_id, file_id)
        result = process_file(content, metadata.filename)

        response = {
            "file_id": file_id,
            "filename": metadata.filename,
            "content_type": result.content_type,
        }

        if result.metadata:
            response["metadata"] = result.metadata

        if result.content_type == "text" and result.text:
            # 文本截断到合理长度
            text = result.text
            if len(text) > 100000:
                text = text[:100000] + f"\n[... truncated, total {len(result.text)} chars ...]"
            response["text"] = text

        elif result.content_type == "image" and result.image_base64:
            response["image"] = {
                "base64": result.image_base64[:100] + "...",  # 摘要（完整 base64 太长）
                "media_type": result.image_media_type,
                "base64_length": len(result.image_base64),
            }
            response["hint"] = "图片已转为 base64。可直接用于多模态 LLM 调用。"

        elif result.error:
            response["error"] = result.error

        return response

    except FileNotFoundError:
        return {"error": f"File not found: {file_id}"}
    except Exception as e:
        logger.error(f"process_file_content error: {e}")
        return {"error": str(e)}
